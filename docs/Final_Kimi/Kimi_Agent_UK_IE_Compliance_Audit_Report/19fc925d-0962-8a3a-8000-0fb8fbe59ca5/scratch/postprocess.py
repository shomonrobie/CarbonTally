#!/usr/bin/env python3
"""Presentation-only house-style post-processing for the CarbonTally hardening plan DOCX.
Zero content changes: geometry, colour palette, table borders/widths, header/footer only."""
import copy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "out/final.footnote.docx"
DST = "out/final.styled.docx"

HEADING = RGBColor(0x2F, 0x5D, 0x4A)
BODY = RGBColor(0x33, 0x33, 0x33)
MUTED = RGBColor(0x6E, 0x8B, 0x7B)
HDR_FILL = "E3EDE7"
BORDER = "9FB8AC"
TITLE = "CarbonTally v1.0 \u2014 Production Hardening Plan"

doc = Document(SRC)

# ---------- 1. Page geometry: A4 landscape, 0.7" margins ----------
for sec in doc.sections:
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Inches(11.69)
    sec.page_height = Inches(8.27)
    sec.left_margin = sec.right_margin = Inches(0.7)
    sec.top_margin = sec.bottom_margin = Inches(0.7)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)

CONTENT_W = Inches(11.69 - 1.4)  # EMU

# ---------- 2. Palette: styles + run-level ----------
def set_style_color(name, rgb):
    try:
        st = doc.styles[name]
        st.font.color.rgb = rgb
    except KeyError:
        pass

for nm in ("Heading 1", "Heading 2", "Heading 3", "Heading 4", "Title"):
    set_style_color(nm, HEADING)
for nm in ("Normal", "Body Text", "First Paragraph", "Compact"):
    set_style_color(nm, BODY)

def color_runs(par, rgb):
    for r in par.runs:
        r.font.color.rgb = rgb

for p in doc.paragraphs:
    if p.style.name.startswith(("Heading", "Title")):
        color_runs(p, HEADING)
    else:
        color_runs(p, BODY)

# ---------- 3. Tables ----------
def set_cell_shading(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexfill)
    tcPr.append(shd)

def table_borders(tbl):
    tblPr = tbl._tbl.tblPr
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), BORDER)
        borders.append(el)
    tblPr.append(borders)

def set_table_width(tbl):
    tblPr = tbl._tbl.tblPr
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), str(int(CONTENT_W / 635)))  # EMU -> twips
    w.set(qn("w:type"), "dxa")
    tblPr.append(w)
    for old in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old)
    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "autofit")
    tblPr.append(lay)

def header_repeat(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)

for tbl in doc.tables:
    table_borders(tbl)
    set_table_width(tbl)
    for ri, row in enumerate(tbl.rows):
        if ri == 0:
            header_repeat(row)
        for cell in row.cells:
            if ri == 0:
                set_cell_shading(cell, HDR_FILL)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = BODY
                    if ri == 0:
                        r.font.bold = True
                        r.font.color.rgb = HEADING

# ---------- 4. Header (right-aligned title + rule) ----------
def add_bottom_rule(par, color=BORDER):
    pPr = par._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_field(par, instr, rgb, size=9):
    def mk_run():
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        c = OxmlElement("w:color"); c.set(qn("w:val"), "%02X%02X%02X" % (rgb[0], rgb[1], rgb[2]))
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(size * 2))
        rPr.append(c); rPr.append(sz)
        r.append(rPr)
        return r
    r1 = mk_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin"); r1.append(f1)
    r2 = mk_run()
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = f" {instr} "; r2.append(it)
    r3 = mk_run()
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end"); r3.append(f3)
    for r in (r1, r2, r3):
        par._p.append(r)

for sec in doc.sections:
    hdr = sec.header
    hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(TITLE)
    run.font.color.rgb = MUTED
    run.font.size = Pt(9)
    add_bottom_rule(hp)

    ftr = sec.footer
    ftr.is_linked_to_previous = False
    fp = ftr.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    def txt(t):
        r = fp.add_run(t)
        r.font.color.rgb = MUTED
        r.font.size = Pt(9)
    txt("Page ")
    add_field(fp, "PAGE", (0x6E, 0x8B, 0x7B))
    txt(" of ")
    add_field(fp, "NUMPAGES", (0x6E, 0x8B, 0x7B))

# ---------- 5. Glyph font fallback (presentation only; characters untouched) ----------
# U+2610/U+2714 render blank under the base font in LibreOffice; U+1F534 needs an emoji font.
# Assign glyph runs to installed symbol fonts; Word substitutes Segoe UI Symbol/Emoji if absent.
GLYPH_FONT = {"\u2610": "Noto Sans Symbols2", "\u2714": "Noto Sans Symbols2",
              "\U0001F534": "Unifont Upper"}
GLYPHS = set(GLYPH_FONT)

def set_run_fonts(r_el, font):
    rPr = r_el.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r_el.insert(0, rPr)
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.insert(0, rf)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(attr), font)

def split_glyph_runs(par):
    for r in list(par.runs):
        text = r.text
        if not text or not any(c in GLYPHS for c in text):
            continue
        # split text into maximal segments of glyph vs non-glyph
        segs, cur, cur_is_g = [], "", None
        for ch in text:
            g = ch in GLYPHS
            if cur_is_g is None or g == cur_is_g:
                cur += ch; cur_is_g = g
            else:
                segs.append((cur, cur_is_g)); cur, cur_is_g = ch, g
        segs.append((cur, cur_is_g))
        orig = r._r
        anchor = orig
        for seg_text, is_g in segs:
            new_r = copy.deepcopy(orig)  # always copy the ORIGINAL run properties
            for t in new_r.findall(qn("w:t")):
                new_r.remove(t)
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = seg_text
            new_r.append(t)
            if is_g:
                set_run_fonts(new_r, GLYPH_FONT[seg_text.strip()[0] if seg_text.strip() else seg_text[0]])
            anchor.addnext(new_r)
            anchor = new_r
        orig.getparent().remove(orig)

def all_paragraphs(document):
    for p in document.paragraphs:
        yield p
    for t in document.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

for p in all_paragraphs(doc):
    split_glyph_runs(p)

doc.save(DST)
print("saved", DST)
