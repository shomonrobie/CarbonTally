#!/usr/bin/env python3
"""Post-process md2docx output to CarbonTally house style (presentation only)."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "/mnt/agents/output/scratch_mm/out/work.footnote.docx"
DST = "/mnt/agents/output/CarbonTally_Migration_Manifest.docx"

HEAD = RGBColor(0x2F, 0x5D, 0x4A)
BODY = RGBColor(0x33, 0x33, 0x33)
HF   = RGBColor(0x6E, 0x8B, 0x7B)
BORDER = "9FB8AC"
SHADE  = "E3EDE7"

doc = Document(SRC)

# --- Page setup: A4 landscape, 0.7" margins ---
for sec in doc.sections:
    sec.page_width  = Inches(11.69)
    sec.page_height = Inches(8.27)
    sec.left_margin = sec.right_margin = Inches(0.7)
    sec.top_margin = sec.bottom_margin = Inches(0.7)

# --- Base styles ---
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = BODY
normal.paragraph_format.space_after = Pt(6)

for hname, size in [("Heading 1", 17), ("Heading 2", 13.5), ("Heading 3", 11.5), ("Heading 4", 10.5)]:
    try:
        st = doc.styles[hname]
    except KeyError:
        continue
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = HEAD
    st.paragraph_format.keep_with_next = True

# Body paragraphs: ensure body colour on runs without explicit colour
for p in doc.paragraphs:
    if p.style.name.startswith("Heading") or p.style.name in ("Title",):
        for r in p.runs:
            r.font.color.rgb = HEAD
    else:
        for r in p.runs:
            r.font.color.rgb = BODY
            if r.font.size is None:
                r.font.size = Pt(10.5)

# --- Tables: full width, borders, header row shading/bold/repeat, 9.5pt cells ---
def set_cell_shading(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

for tbl in doc.tables:
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    tblPr = tbl._tbl.tblPr
    # layout autofit + 100% width
    for el in tblPr.findall(qn("w:tblLayout")): tblPr.remove(el)
    lay = OxmlElement("w:tblLayout"); lay.set(qn("w:type"), "autofit"); tblPr.append(lay)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000"); tblW.set(qn("w:type"), "pct")
    # borders
    for el in tblPr.findall(qn("w:tblBorders")): tblPr.remove(el)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), BORDER)
        borders.append(e)
    tblPr.append(borders)

    for ri, row in enumerate(tbl.rows):
        if ri == 0:
            trPr = row._tr.get_or_add_trPr()
            tblHeader = OxmlElement("w:tblHeader"); tblHeader.set(qn("w:val"), "true")
            trPr.append(tblHeader)
        for cell in row.cells:
            if ri == 0:
                set_cell_shading(cell, SHADE)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
                    if ri == 0:
                        r.font.bold = True
                        r.font.color.rgb = HEAD
                    else:
                        r.font.color.rgb = BODY

# --- Header: right-aligned title + bottom rule ---
sec = doc.sections[0]
hp = sec.header.paragraphs[0]
hp.text = ""
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("CarbonTally — Migration Manifest (RC1)")
hr.font.name = "Calibri"; hr.font.size = Pt(9); hr.font.color.rgb = HF; hr.font.italic = True
pPr = hp._p.get_or_add_pPr()
pbdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), BORDER)
pbdr.append(bottom); pPr.append(pbdr)

# --- Footer: centred Page X of Y ---
fp = sec.footer.paragraphs[0]
fp.text = ""
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_field(par, instr):
    r1 = par.add_run(); f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin"); r1._r.append(f1)
    r2 = par.add_run(); it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = f" {instr} "; r2._r.append(it)
    r3 = par.add_run(); f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end"); r3._r.append(f3)
    return (r1, r2, r3)

def style_run(r):
    r.font.name = "Calibri"; r.font.size = Pt(9); r.font.color.rgb = HF

style_run(fp.add_run("Page "))
for r in add_field(fp, "PAGE"): style_run(r)
style_run(fp.add_run(" of "))
for r in add_field(fp, "NUMPAGES"): style_run(r)

doc.save(DST)
print("saved", DST)
